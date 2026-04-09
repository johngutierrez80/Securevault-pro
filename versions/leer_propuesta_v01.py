from docx import Document
from pathlib import Path

docx_path = Path(r"c:\Users\John\Downloads\securevault_pro_full\versions\v0.1-inicial\Propuesta_Inicio_Desarrollo_v0.1.docx")
out_path = Path(r"c:\Users\John\Downloads\securevault_pro_full\versions\v0.1-inicial\Propuesta_Inicio_Desarrollo_v0.1.txt")

doc = Document(docx_path)
lines = []
for p in doc.paragraphs:
    text = p.text.strip()
    if text:
        lines.append(text)

for table in doc.tables:
    lines.append("")
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        lines.append(" | ".join(cells))

out_path.write_text("\n".join(lines), encoding="utf-8")
print(str(out_path))
