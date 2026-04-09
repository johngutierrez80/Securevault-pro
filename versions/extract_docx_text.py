from docx import Document
from pathlib import Path

src = Path(r"c:\Users\John\Downloads\securevault_pro_full\Proyecto.docx")
out = Path(r"c:\Users\John\Downloads\securevault_pro_full\versions\Proyecto_extraido.txt")

doc = Document(src)
lines = []
for p in doc.paragraphs:
    t = p.text.strip()
    if t:
        lines.append(t)

for table in doc.tables:
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        lines.append(" | ".join(cells))

out.write_text("\n".join(lines), encoding="utf-8")
print(f"Extraido: {out}")
