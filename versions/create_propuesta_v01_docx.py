from datetime import date
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

out_path = r"c:\Users\John\Downloads\securevault_pro_full\versions\v0.1-inicial\Propuesta_Inicio_Desarrollo_v0.1.docx"

doc = Document()

# Portada
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = title.add_run("SecureVault Pro\nPropuesta de Inicio de Desarrollo (v0.1)")
run.bold = True
run.font.size = None

subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle.add_run("Documento adaptado a partir de Proyecto.docx y alineado con la fase inicial del proyecto.")

meta = doc.add_paragraph()
meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
meta.add_run(f"Fecha: {date.today().isoformat()}\nVersión: v0.1\nEstado: Propuesta aprobable para arranque")

doc.add_page_break()

# 1. Resumen ejecutivo
doc.add_heading("1. Resumen Ejecutivo", level=1)
doc.add_paragraph(
    "Este documento define la propuesta de inicio de desarrollo de SecureVault Pro en su versión v0.1. "
    "La propuesta prioriza una implementación funcional y demostrable de arquitectura de microservicios con "
    "autenticación centralizada, gestión de secretos y despliegue local con Docker Compose. "
    "Se delimita el alcance de la fase inicial para asegurar una base sólida antes de incorporar capacidades avanzadas "
    "de DevSecOps, observabilidad y endurecimiento productivo."
)

# 2. Contexto y justificación
doc.add_heading("2. Contexto y Justificación", level=1)
doc.add_paragraph(
    "El documento Proyecto.docx plantea una visión completa de ciberseguridad y DevSecOps. "
    "Para una ejecución controlada, esta adaptación concentra el inicio del desarrollo en un núcleo funcional mínimo "
    "viable, manteniendo alineación con buenas prácticas de seguridad: autenticación con JWT, hash de contraseñas, "
    "cifrado de secretos y separación de responsabilidades por servicio."
)

# 3. Objetivos
doc.add_heading("3. Objetivos", level=1)
doc.add_paragraph("Objetivo general:", style=None).runs[0].bold = True
doc.add_paragraph(
    "Implementar una base funcional de SecureVault Pro (v0.1) que permita registrar usuarios, autenticar accesos "
    "y operar un CRUD de secretos cifrados en arquitectura de microservicios."
)
doc.add_paragraph("Objetivos específicos:").runs[0].bold = True
for item in [
    "Definir y desplegar servicios auth-service y vault-service con FastAPI.",
    "Integrar PostgreSQL para persistencia de usuarios y secretos.",
    "Exponer un frontend básico para login y gestión inicial de secretos.",
    "Orquestar el entorno local con Docker Compose y Nginx como gateway.",
    "Documentar manuales base de usuario, técnico, operación, seguridad y pruebas.",
]:
    doc.add_paragraph(item, style="List Bullet")

# 4. Alcance v0.1
doc.add_heading("4. Alcance de la Fase Inicial (v0.1)", level=1)
doc.add_paragraph("Incluye:").runs[0].bold = True
for item in [
    "Registro y login de usuario con emisión de JWT.",
    "Gestión de secretos (crear, listar, editar, eliminar) por usuario autenticado.",
    "Hash de contraseñas y cifrado de secretos en backend.",
    "Despliegue local con Docker Compose.",
    "Documentación base de entrega.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph("No incluye en esta fase:").runs[0].bold = True
for item in [
    "Pipeline CI/CD completo de publicación y despliegue.",
    "Observabilidad integral (Prometheus, Loki, Grafana, Falco).",
    "Automatización IaC para ambientes productivos.",
    "Hardening avanzado de contenedores y escaneo continuo.",
]:
    doc.add_paragraph(item, style="List Bullet")

# 5. Arquitectura base
doc.add_heading("5. Arquitectura Base Propuesta", level=1)
doc.add_paragraph(
    "La versión v0.1 conserva una arquitectura de microservicios ligera para acelerar validación funcional:"
)
for item in [
    "Auth Service (FastAPI, puerto 8001): registro/login, emisión de tokens JWT.",
    "Vault Service (FastAPI, puerto 8002): CRUD de secretos por usuario autenticado.",
    "PostgreSQL: persistencia de usuarios y secretos.",
    "Redis: soporte para control de tasa en Vault.",
    "Gateway/Frontend: acceso web unificado y consumo de API.",
]:
    doc.add_paragraph(item, style="List Bullet")

# 6. Entregables
doc.add_heading("6. Entregables de Inicio", level=1)
for item in [
    "Repositorio organizado por servicios y documentación.",
    "docker-compose.yml funcional para despliegue local.",
    "Servicios auth-service y vault-service operativos.",
    "Frontend básico operativo para flujo de autenticación y bóveda.",
    "Documentación inicial en carpeta docs.",
    "Snapshot de versión inicial en versions/v0.1-inicial.",
]:
    doc.add_paragraph(item, style="List Bullet")

# 7. Plan de implementación
doc.add_heading("7. Plan de Implementación", level=1)

table = doc.add_table(rows=1, cols=4)
header = table.rows[0].cells
header[0].text = "Fase"
header[1].text = "Objetivo"
header[2].text = "Resultado esperado"
header[3].text = "Estado"

rows = [
    ("Fase 1: Inicio (v0.1)", "Base funcional mínima", "Auth + Vault + DB + Frontend básico + documentación", "Completada"),
    ("Fase 2: Calidad y Seguridad", "Elevar calidad técnica", "Tests, rate limiting reforzado, SPA y modelado de amenazas", "Planificada/Parcial"),
    ("Fase 3: DevSecOps", "Preparar operación continua", "CI/CD completo, monitoreo, seguridad runtime e IaC", "Planificada"),
]
for r in rows:
    cells = table.add_row().cells
    cells[0].text, cells[1].text, cells[2].text, cells[3].text = r

# 8. Criterios de aceptación
doc.add_heading("8. Criterios de Aceptación para Cierre de v0.1", level=1)
for item in [
    "El entorno levanta con Docker Compose sin errores críticos.",
    "El flujo registro/login entrega token válido.",
    "Los endpoints de bóveda requieren autenticación y operan CRUD.",
    "Las credenciales de usuario no se almacenan en texto plano.",
    "La documentación base permite ejecutar y demostrar el sistema.",
]:
    doc.add_paragraph(item, style="List Bullet")

# 9. Riesgos y mitigación
doc.add_heading("9. Riesgos Iniciales y Mitigaciones", level=1)
for item in [
    "Riesgo: deuda técnica por aceleración inicial. Mitigación: fase 2 enfocada en pruebas y refactor controlado.",
    "Riesgo: configuración insegura por entorno local. Mitigación: uso de .env.example y segregación de secretos.",
    "Riesgo: cobertura de seguridad limitada en inicio. Mitigación: roadmap explícito hacia DevSecOps y monitoreo.",
]:
    doc.add_paragraph(item, style="List Bullet")

# 10. Aprobación
doc.add_heading("10. Declaración de Propuesta", level=1)
doc.add_paragraph(
    "Se propone adoptar esta versión v0.1 como línea base formal de inicio de desarrollo, "
    "manteniendo trazabilidad hacia fases posteriores de calidad, seguridad y operación DevSecOps."
)

doc.save(out_path)
print(f"Documento generado: {out_path}")
